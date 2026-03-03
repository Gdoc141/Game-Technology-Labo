"""
Converteert VERSLAG_FASE1.md naar een opgemaakt Word-document (.docx)
Gebruik: python md_to_word.py
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE  = r"l:\1.Documents2\School\GameTech\Game-Technology-Labo\Opdracht_1\Lasertag_opdracht1\VERSLAG_FASE1.md"
OUT_FILE = r"l:\1.Documents2\School\GameTech\Game-Technology-Labo\Opdracht_1\Lasertag_opdracht1\VERSLAG_FASE1.docx"


# ────────────────────────────────────────────────
# Hulpfuncties
# ────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill_hex="404040"):
    """Grijze achtergrond voor een tabelcel."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def add_code_block(doc, code_text):
    """Voegt een lichtgrijs code-blok toe."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)

    # Achtergrondkleur via XML shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)

    run = para.add_run(code_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    return para


def add_inline_formatting(para, text):
    """Verwerkt **vet**, *cursief* en `code` inline in een paragraaf."""
    pattern = re.compile(r'(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*)')
    last = 0
    for m in pattern.finditer(text):
        # Tekst vóór match
        if m.start() > last:
            run = para.add_run(text[last:m.start()])
            run.font.size = Pt(11)

        full = m.group(0)
        if full.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(11)
        elif full.startswith("`"):
            run = para.add_run(m.group(3))
            run.font.name  = "Courier New"
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif full.startswith("*"):
            run = para.add_run(m.group(4))
            run.italic = True
            run.font.size = Pt(11)
        last = m.end()

    # Resterende tekst
    if last < len(text):
        run = para.add_run(text[last:])
        run.font.size = Pt(11)


def parse_table(doc, lines):
    """Verwerkt een markdown tabel en voegt hem toe aan het document."""
    rows = []
    for line in lines:
        line = line.strip()
        if re.match(r'^\|[-| :]+\|$', line):
            continue
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = row.cells[c_idx]
            p    = cell.paragraphs[0]
            p.clear()

            plain = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
            plain = re.sub(r'`(.+?)`', r'\1', plain)

            run = p.add_run(plain)
            run.font.size = Pt(10)

            if r_idx == 0:  # Header rij
                run.bold = True
                shade_cell(cell, "D0E4F5")
            else:
                run.bold = bool(re.search(r'\*\*', cell_text))

    doc.add_paragraph()  # Witruimte na tabel


# ────────────────────────────────────────────────
# Hoofd-conversie
# ────────────────────────────────────────────────

def convert(md_path, out_path):
    doc = Document()

    # Paginamarges instellen
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

    # Standaard lettertype instellen
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        raw  = lines[i]
        line = raw.rstrip("\n")

        # ── Lege regel
        if line.strip() == "":
            i += 1
            continue

        # ── Horizontale lijn (---)
        if re.match(r'^---+$', line.strip()):
            para = doc.add_paragraph()
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "999999")
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Codeblok (``` ... ```)
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1  # sla sluitende ``` over
            code_text = "\n".join(code_lines)
            add_code_block(doc, code_text)
            continue

        # ── Markdown tabel
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            parse_table(doc, table_lines)
            continue

        # ── Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text  = heading_match.group(2)
            # Verwijder markdown links [text](url) → text
            text  = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            # Verwijder nummering-ankertekst (## 1. Titel)
            text  = re.sub(r'^\d+\.\s+', '', text)

            word_level = min(level, 4)
            heading_names = {1: "Heading 1", 2: "Heading 2",
                             3: "Heading 3", 4: "Heading 4"}
            para = doc.add_paragraph(style=heading_names[word_level])
            run  = para.add_run(text)

            sizes = {1: 18, 2: 15, 3: 13, 4: 12}
            colors = {
                1: (0x1F, 0x39, 0x64),
                2: (0x2E, 0x74, 0xB5),
                3: (0x25, 0x63, 0x99),
                4: (0x40, 0x40, 0x40),
            }
            run.font.size  = Pt(sizes[word_level])
            run.font.color.rgb = RGBColor(*colors[word_level])
            run.font.bold  = True
            i += 1
            continue

        # ── Bullet lijst (- of *)
        if re.match(r'^[\-\*]\s+', line):
            text = re.sub(r'^[\-\*]\s+', '', line)
            para = doc.add_paragraph(style="List Bullet")
            add_inline_formatting(para, text)
            i += 1
            continue

        # ── Genummerde lijst
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            para = doc.add_paragraph(style="List Number")
            add_inline_formatting(para, text)
            i += 1
            continue

        # ── Citaat (>)
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.0)
            pPr  = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            lft  = OxmlElement("w:left")
            lft.set(qn("w:val"),   "single")
            lft.set(qn("w:sz"),    "12")
            lft.set(qn("w:space"), "4")
            lft.set(qn("w:color"), "2E74B5")
            pBdr.append(lft)
            pPr.append(pBdr)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            i += 1
            continue

        # ── Gewone paragraaf
        # Verwijder markdown link-syntax [text](url) → text
        clean = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', line)
        para  = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        add_inline_formatting(para, clean)
        i += 1

    doc.save(out_path)
    print(f"✓ Opgeslagen: {out_path}")


if __name__ == "__main__":
    convert(MD_FILE, OUT_FILE)
